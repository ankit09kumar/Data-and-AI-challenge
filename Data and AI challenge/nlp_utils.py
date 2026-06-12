import re
import pdfplumber
import docx
from datetime import datetime
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import json

# Comprehensive database of standard developer and business skills
COMMON_SKILLS = [
    # Languages
    'python', 'javascript', 'java', 'c++', 'c#', 'c', 'ruby', 'go', 'golang', 'rust', 'php', 'swift', 'kotlin', 'typescript', 'scala', 'r', 'matlab', 'perl', 'shell', 'bash',
    # Frontend
    'html', 'css', 'sass', 'bootstrap', 'tailwind', 'react', 'angular', 'vue', 'nextjs', 'jquery', 'svelte',
    # Backend / Web frameworks
    'flask', 'django', 'fastapi', 'node.js', 'node', 'express', 'spring boot', 'laravel', 'asp.net', '.net', 'rails',
    # Databases
    'sql', 'mysql', 'postgresql', 'sqlite', 'mongodb', 'redis', 'elasticsearch', 'cassandra', 'dynamodb', 'oracle', 'mariadb',
    # AI / ML / Data Science
    'machine learning', 'deep learning', 'nlp', 'natural language processing', 'computer vision', 'data science', 'pandas', 'numpy', 'scikit-learn', 'tensorflow', 'pytorch', 'keras', 'opencv', 'spacy', 'nltk', 'llm', 'bert', 'gpt', 'reinforcement learning', 'data analysis', 'data visualization', 'tableau', 'power bi',
    # DevOps / Cloud
    'aws', 'azure', 'gcp', 'google cloud', 'docker', 'kubernetes', 'jenkins', 'git', 'github', 'gitlab', 'ci/cd', 'terraform', 'ansible', 'linux', 'unix',
    # Architecture / Methodologies
    'rest api', 'graphql', 'microservices', 'mvc', 'agile', 'scrum', 'oop', 'object-oriented programming', 'system design',
    # Soft & Management Skills
    'project management', 'communication', 'leadership', 'teamwork', 'problem solving', 'critical thinking', 'collaboration'
]

def extract_text_from_pdf(file_path):
    """Extracts raw text from a PDF document using pdfplumber."""
    text = ""
    try:
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                content = page.extract_text()
                if content:
                    text += content + "\n"
    except Exception as e:
        print(f"Error reading PDF {file_path}: {e}")
    return text

def extract_text_from_docx(file_path):
    """Extracts raw text from a DOCX document using python-docx."""
    text = ""
    try:
        doc = docx.Document(file_path)
        for para in doc.paragraphs:
            text += para.text + "\n"
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
    except Exception as e:
        print(f"Error reading DOCX {file_path}: {e}")
    return text

def extract_contact_info(text):
    """Extracts Name, Email, and Phone number using heuristics and regular expressions."""
    # Find email
    email_pattern = r'[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+'
    email_match = re.search(email_pattern, text)
    email = email_match.group(0) if email_match else ""

    # Find phone (supports formats like: +1-234-567-8900, (123) 456-7890, 1234567890, 123-456-7890)
    phone_pattern = r'(?:\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}\b'
    phone_match = re.search(phone_pattern, text)
    phone = phone_match.group(0) if phone_match else ""

    # Extract name (heuristic: first non-empty lines, filtering out contact coordinates)
    name = "Unknown Candidate"
    lines = [line.strip() for line in text.split('\n') if line.strip()]
    for line in lines[:5]:
        # Clean line
        cleaned = re.sub(r'[^a-zA-Z\s]', '', line).strip()
        # Avoid lines with email domains or excessive numbers
        if (
            '@' not in line 
            and '.com' not in line.lower()
            and len(cleaned.split()) >= 2 
            and len(cleaned.split()) <= 4
            and not any(keyword in line.lower() for keyword in ['curriculum', 'resume', 'cv', 'experience', 'summary'])
        ):
            name = cleaned
            break
            
    return name, email, phone

def extract_skills(text):
    """Matches text against common skills dictionary with boundary-aware regex."""
    text_lower = text.lower()
    found_skills = set()
    
    for skill in COMMON_SKILLS:
        # Create custom boundary checks for special chars (+, #, etc.)
        if skill == 'c++':
            # Matches c++ surrounded by whitespace, punctuation, or boundaries
            pattern = r'(?:^|[\s,;:\(\)])c\+\+(?:$|[\s,;:\.\(\)])'
        elif skill == 'c#':
            pattern = r'(?:^|[\s,;:\(\)])c\#(?:$|[\s,;:\.\(\)])'
        elif skill == '.net':
            pattern = r'(?:^|[\s,;:\(\)])\.net(?:$|[\s,;:\.\(\)])'
        elif len(skill) <= 2:
            # Short skills (e.g. "c", "r", "go") need strict word boundaries
            pattern = r'\b' + re.escape(skill) + r'\b'
        else:
            # General skill matching
            pattern = r'\b' + re.escape(skill) + r'\b'
            
        if re.search(pattern, text_lower):
            found_skills.add(skill)
            
    return list(found_skills)

def extract_experience_years(text):
    """Calculates experience using explicit years statements and work durations."""
    text_lower = text.lower()
    
    # 1. Look for explicit assertions: e.g. "5 years of experience", "experience: 3 yrs"
    patterns = [
        r'(\d+(?:\.\d+)?)\s*\+?\s*years?\s*of\s*(?:work\s+)?experience\b',
        r'(\d+(?:\.\d+)?)\s*\+?\s*years?\s*(?:of\s+)?expertise\b',
        r'(\d+(?:\.\d+)?)\s*\+?\s*yrs?\s*(?:of\s+)?experience\b',
        r'experience\s*[:\-]\s*(\d+(?:\.\d+)?)\s*years?\b',
        r'(\d+(?:\.\d+)?)\s*years?\s*experience\b',
        r'(\d+(?:\.\d+)?)\s*\+?\s*years?\s*in\b'
    ]
    
    explicit_years = []
    for pattern in patterns:
        matches = re.findall(pattern, text_lower)
        for m in matches:
            try:
                val = float(m)
                if 0.5 <= val <= 40.0:
                    explicit_years.append(val)
            except ValueError:
                pass
                
    # 2. Extract date ranges and aggregate duration (e.g., "2018 - 2022", "2019 to Present")
    current_year = datetime.now().year
    date_patterns = [
        r'\b(19\d{2}|20\d{2})\s*[-–—–to]+\s*(19\d{2}|20\d{2})\b',
        r'\b(19\d{2}|20\d{2})\s*[-–—–to]+\s*(present|current|now|date)\b'
    ]
    
    range_years = 0.0
    # Capture all date range segments
    for pattern in date_patterns:
        matches = re.findall(pattern, text_lower)
        for m in matches:
            try:
                start = int(m[0])
                if m[1] in ['present', 'current', 'now', 'date']:
                    end = current_year
                else:
                    end = int(m[1])
                
                duration = end - start
                if 0 < duration <= 40:
                    range_years += duration
            except ValueError:
                pass
                
    # Heuristic combination:
    # If explicitly declared years exist, they are highly reliable. Take the max.
    # Otherwise, fall back to aggregate ranges.
    if explicit_years:
        return max(explicit_years)
    elif range_years > 0.0:
        # Prevent double counting overlapping ranges by capping aggregate at 25
        return min(range_years, 30.0)
    
    return 0.0

def calculate_semantic_similarity(candidate_text, job_desc_text):
    """Calculates TF-IDF Cosine Similarity between resume text and job requirements."""
    if not candidate_text.strip() or not job_desc_text.strip():
        return 0.0
    try:
        vectorizer = TfidfVectorizer(stop_words='english')
        tfidf = vectorizer.fit_transform([candidate_text, job_desc_text])
        similarity = cosine_similarity(tfidf[0:1], tfidf[1:2])[0][0]
        return float(similarity * 100)
    except Exception as e:
        print(f"Error calculating similarity: {e}")
        return 0.0

def match_candidate_to_job(candidate, job):
    """
    Ranks a candidate against a job description.
    Returns: (match_score, skill_score, experience_score, semantic_score, match_details_json)
    
    Weights:
    - Skill Match Score: 40%
    - Experience Match Score: 30%
    - Semantic text similarity: 30%
    """
    # 1. Skill Match Score (Jaccard-like overlap coefficient relative to required list)
    job_skills_list = [s.strip().lower() for s in job.skills_required.split(',') if s.strip()]
    if not job_skills_list:
        skill_score = 100.0
        matched_skills = []
        missing_skills = []
    else:
        candidate_skills_list = [s.strip().lower() for s in (candidate.skills_extracted or "").split(',') if s.strip()]
        matched_skills = list(set(job_skills_list).intersection(set(candidate_skills_list)))
        missing_skills = list(set(job_skills_list).difference(set(candidate_skills_list)))
        skill_score = (len(matched_skills) / len(job_skills_list)) * 100.0
        
    # 2. Experience Match Score (30%)
    req_exp = float(job.experience_years_required or 0)
    cand_exp = float(candidate.experience_years or 0.0)
    
    if req_exp <= 0:
        experience_score = 100.0
    elif cand_exp >= req_exp:
        experience_score = 100.0
    else:
        # Partial credit for experience
        experience_score = (cand_exp / req_exp) * 100.0
        
    # 3. Semantic Similarity Score (30%)
    semantic_score = calculate_semantic_similarity(candidate.parsed_text or "", job.description or "")
    
    # Combined Weighted Score
    match_score = (0.4 * skill_score) + (0.3 * experience_score) + (0.3 * semantic_score)
    
    # Generate Recommendation Statement
    status_msg = ""
    if match_score >= 80:
        status_msg = "Highly Recommended: Excellent skill fit and experience matches or exceeds job requirements."
    elif match_score >= 60:
        status_msg = "Recommended: Strong candidate with a solid skill base, matching some core requirements."
    elif match_score >= 40:
        status_msg = "Potential Match: Meets basic qualifications but shows gaps in required skills or experience."
    else:
        status_msg = "Low Match: Significant divergence from required experience levels or technical skill set."

    details = {
        'matched_skills': matched_skills,
        'missing_skills': missing_skills,
        'experience_required': req_exp,
        'experience_candidate': cand_exp,
        'recommendation_rationale': status_msg
    }
    
    return (
        round(match_score, 1),
        round(skill_score, 1),
        round(experience_score, 1),
        round(semantic_score, 1),
        json.dumps(details)
    )
